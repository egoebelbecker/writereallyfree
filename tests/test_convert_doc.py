import os
from docx import Document
from convert_doc import add_runs_to_paragraph, create_word_document
from bs4 import BeautifulSoup
import markdown

def test_add_runs_to_paragraph():
    md = "Hello **world** and *everyone* and ***bold-italic***."
    html = markdown.markdown(md)
    soup = BeautifulSoup(html, "html.parser")
    
    # We should have a single <p> element
    p_element = soup.find('p')
    assert p_element is not None
    
    doc = Document()
    p = doc.add_paragraph()
    add_runs_to_paragraph(p, p_element)
    
    # We expect several runs:
    # 1. "Hello "
    # 2. "world" (bold=True)
    # 3. " and "
    # 4. "everyone" (italic=True)
    # 5. " and "
    # 6. "bold-italic" (bold=True, italic=True)
    # 7. "."
    runs = p.runs
    assert len(runs) >= 6
    
    # Find the run containing "world"
    world_run = next(r for r in runs if r.text == "world")
    assert world_run.bold is True
    assert world_run.italic is not True
    
    # Find the run containing "everyone"
    everyone_run = next(r for r in runs if r.text == "everyone")
    assert everyone_run.bold is not True
    assert everyone_run.italic is True
    
    # Find the run containing "bold-italic"
    bi_run = next(r for r in runs if r.text == "bold-italic")
    assert bi_run.bold is True
    assert bi_run.italic is True

def test_horizontal_rule(tmp_path):
    from docx.oxml.ns import qn
    md_content = "Paragraph 1\n\n---\n\nParagraph 2"
    temp_md = tmp_path / "test.md"
    temp_md.write_text(md_content, encoding='utf-8')
    
    create_word_document(str(tmp_path), "test.md")
    
    output_docx = tmp_path / "test.docx"
    assert output_docx.exists()
    
    doc = Document(str(output_docx))
    assert len(doc.paragraphs) == 3
    
    # Check that the second paragraph has bottom border XML settings
    p = doc.paragraphs[1]
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    assert pBdr is not None
    bottom = pBdr.find(qn('w:bottom'))
    assert bottom is not None
    assert bottom.get(qn('w:val')) == 'single'

def test_hyperlink_preservation():
    from docx.oxml.ns import qn
    md = "Please visit [Example Website](https://example.com) now."
    html = markdown.markdown(md)
    soup = BeautifulSoup(html, "html.parser")
    
    p_element = soup.find('p')
    assert p_element is not None
    
    doc = Document()
    p = doc.add_paragraph()
    add_runs_to_paragraph(p, p_element)
    
    # Check that the paragraph contains the w:hyperlink node
    hyperlink_node = p._p.find(qn('w:hyperlink'))
    assert hyperlink_node is not None
    
    # Verify the relationship ID is bound
    r_id = hyperlink_node.get(qn('r:id'))
    assert r_id is not None
    
    # Verify the text inside the hyperlink
    r_node = hyperlink_node.find(qn('w:r'))
    assert r_node is not None
    t_node = r_node.find(qn('w:t'))
    assert t_node is not None
    assert t_node.text == "Example Website"


def test_create_word_document_doublespace(tmp_path):
    md_content = "# Title\n\nParagraph 1\n\nParagraph 2"
    temp_md = tmp_path / "doublespace_test.md"
    temp_md.write_text(md_content, encoding='utf-8')

    create_word_document(str(tmp_path), "doublespace_test.md", doublespace=True)

    output_docx = tmp_path / "doublespace_test.docx"
    assert output_docx.exists()

    doc = Document(str(output_docx))
    assert len(doc.paragraphs) == 3
    for p in doc.paragraphs:
        assert p.paragraph_format.line_spacing == 2.0


def test_create_word_document_paragraph_formatting(tmp_path):
    md_content = "# Title\n\nBody Paragraph"
    temp_md = tmp_path / "format_test.md"
    temp_md.write_text(md_content, encoding='utf-8')

    create_word_document(str(tmp_path), "format_test.md", indent_first_line=True, space_before=True, space_after=True)

    output_docx = tmp_path / "format_test.docx"
    assert output_docx.exists()

    doc = Document(str(output_docx))
    assert len(doc.paragraphs) == 2

    # Heading paragraph (not body)
    heading_p = doc.paragraphs[0]
    assert heading_p.paragraph_format.space_before.pt == 12.0
    assert heading_p.paragraph_format.space_after.pt == 12.0

    # Body paragraph
    body_p = doc.paragraphs[1]
    assert body_p.paragraph_format.first_line_indent.inches == 0.5
    assert body_p.paragraph_format.space_before.pt == 12.0
    assert body_p.paragraph_format.space_after.pt == 12.0




