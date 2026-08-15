"""

This file owes a lot to Craig Wilson over here:
https://www.craigwilson.blog/post/2025/2025-02-20-creatingwordfrommd/

But I've removed his code that traversed links and handled tables and
added support for a few other tags.

Dependencies:
    - markdown
    - beautifulsoup4
    - python-docx
"""

import os
import sys
import re
import markdown
from bs4 import BeautifulSoup
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches, Pt


def change_extension_to_docx(filename):
    """Return a filename with its extension changed to .docx.

    If the filename has an extension, it is replaced with .docx. If it
    does not have an extension, .docx is appended.

    Args:
        filename (str): The input filename.

    Returns:
        str: The filename with a .docx extension.
    """
    base, ext = os.path.splitext(filename)
    if not ext:
        return f"{filename}.docx"
    return f"{base}.docx"

def add_hyperlink_to_paragraph(paragraph, url, anchor_node, bold=False, italic=False):
    """
    Adds a hyperlink with correct style and nested formatting to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink_el = OxmlElement('w:hyperlink')
    hyperlink_el.set(qn('r:id'), r_id)
    
    def traverse_link_node(node, current_bold=False, current_italic=False):
        if node.name is None:
            r_el = OxmlElement('w:r')
            run = docx.text.run.Run(r_el, paragraph)
            run.text = str(node)
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 238)
            if current_bold:
                run.bold = True
            if current_italic:
                run.italic = True
            hyperlink_el.append(r_el)
        else:
            new_bold = current_bold or (node.name in ['strong', 'b'])
            new_italic = current_italic or (node.name in ['em', 'i'])
            for sub in node.children:
                traverse_link_node(sub, new_bold, new_italic)
                
    traverse_link_node(anchor_node, bold, italic)
    paragraph._p.append(hyperlink_el)

def add_runs_to_paragraph(paragraph, parent_element):
    """
    Recursively processes parent_element's children (text nodes and inline formatting tags)
    and adds them as runs to the paragraph with correct bold/italic/link formatting.
    """
    for child in parent_element.children:
        if child.name is None:  # Plain text node
            paragraph.add_run(child)
        elif child.name == 'a':
            url = child.get('href', '')
            add_hyperlink_to_paragraph(paragraph, url, child)
        else:
            is_bold = child.name in ['strong', 'b']
            is_italic = child.name in ['em', 'i']
            
            def traverse(node, bold=False, italic=False):
                if node.name is None:
                    run = paragraph.add_run(node)
                    if bold:
                        run.bold = True
                    if italic:
                        run.italic = True
                elif node.name == 'a':
                    url = node.get('href', '')
                    add_hyperlink_to_paragraph(paragraph, url, node, bold, italic)
                else:
                    new_bold = bold or (node.name in ['strong', 'b'])
                    new_italic = italic or (node.name in ['em', 'i'])
                    for sub_child in node.children:
                        traverse(sub_child, new_bold, new_italic)
            
            traverse(child, is_bold, is_italic)

def add_horizontal_rule(paragraph):
    """
    Adds a horizontal rule (bottom border) to a python-docx paragraph.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_word_document(folder, file_name, doublespace=False, indent_first_line=False, space_before=False, space_after=False):
    input_path = os.path.join(folder, file_name)
    output_name = change_extension_to_docx(file_name)
    output_path = os.path.join(folder, output_name)

    if os.path.exists(input_path):
        with open(input_path, 'r', encoding='utf-8') as f:
            md = f.read()

            doc = Document()
            try:
                style_normal = doc.styles['Normal']
                if doublespace:
                    style_normal.paragraph_format.line_spacing = 2.0
                if indent_first_line:
                    style_normal.paragraph_format.first_line_indent = Inches(0.5)
                if space_before:
                    style_normal.paragraph_format.space_before = Pt(12)
                if space_after:
                    style_normal.paragraph_format.space_after = Pt(12)
            except Exception:
                pass

            html = markdown.markdown(md)
            soup = BeautifulSoup(html, "html.parser")

            def format_p(p, is_body=False):
                if p is None:
                    return
                if doublespace:
                    p.paragraph_format.line_spacing = 2.0
                if indent_first_line and is_body:
                    p.paragraph_format.first_line_indent = Inches(0.5)
                if space_before:
                    p.paragraph_format.space_before = Pt(12)
                if space_after:
                    p.paragraph_format.space_after = Pt(12)

            for element in soup.contents:
                p = None
                if element.name and len(element.name) == 2 and element.name[0] == 'h' and element.name[1].isdigit():
                    level = int(element.name[1])
                    p = doc.add_heading('', level=level)
                    add_runs_to_paragraph(p, element)
                    format_p(p, is_body=False)
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    add_runs_to_paragraph(p, element)
                    format_p(p, is_body=True)
                elif element.name in ['ul', 'ol']:
                    style = 'List Bullet' if element.name == 'ul' else 'List Number'
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(style=style)
                        add_runs_to_paragraph(p, li)
                        format_p(p, is_body=False)
                elif element.name == 'blockquote':
                    p = doc.add_paragraph(style='Intense Quote')
                    add_runs_to_paragraph(p, element)
                    format_p(p, is_body=False)
                elif element.name == 'hr':
                    p = doc.add_paragraph()
                    add_horizontal_rule(p)
                    format_p(p, is_body=False)

            doc.save(output_path)

    else:
        print(f"WARNING: File not found - {input_path}")
    print(f"Document saved to: {output_path}")
